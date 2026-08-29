"""Build the fillable service-directory Word document.

    python scripts/build_service_doc.py

A landscape table with the rows already in the CSV pre-filled, and blank rows
underneath. Whatever comes back gets typed into services.csv, so the column
order here matches the CSV exactly.
"""

from __future__ import annotations

import csv
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CSV = ROOT / "data" / "services" / "services.csv"
OUT = ROOT / "docs" / "service_directory_to_fill.docx"

PLUM = RGBColor(0x5B, 0x23, 0x40)
TEAL = RGBColor(0x0E, 0x7A, 0x86)
RUST = RGBColor(0xC0, 0x4B, 0x2F)
MUTED = RGBColor(0x70, 0x68, 0x60)

#: Column, printed heading, width in inches. Narrow columns first so the wide
#: free-text ones get the remaining space.
COLUMNS = [
    ("service_id", "ID", 0.8),
    ("name", "Name", 1.5),
    ("service_type", "Type", 0.9),
    ("routes", "Routes it serves", 1.5),
    ("contact_type", "Channel", 0.9),
    ("contact", "Number / address", 1.4),
    ("coverage", "Coverage", 0.9),
    ("is_free", "Free?", 0.5),
    ("anonymous_ok", "Anon?", 0.5),
    ("opening_hours", "Hours", 0.9),
    ("eligibility", "Who can use it", 1.2),
    ("what_they_do", "One sentence, in her words", 2.0),
    ("source", "Source", 1.1),
    ("verified_by", "Checked by", 0.8),
    ("verified_at", "Date", 0.7),
    ("status", "Status", 0.7),
]

BLANK_ROWS = 14


def para(doc, text, size=10.5, bold=False, color=None, space=6, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return p


def main() -> int:
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Inches(0.4))

    para(doc, "Trusted Aunti — verified service directory", 20, True, PLUM, 2)
    para(doc, "Contacts are read from this table and never written by the model. "
              "Nothing here reaches a girl until Status says verified.",
         11, False, MUTED, 14)

    # --- how to fill it ------------------------------------------------------
    para(doc, "Four things that matter more than filling every cell",
         13, True, TEAL, 6)
    for line in [
        "A text channel beats a voice one. A girl on a shared phone, in a room "
        "with family, often cannot make a call — so an SMS or WhatsApp number is "
        "worth more than another hotline.",
        "Leave a cell blank rather than guessing. Blank means unknown and the "
        "system will not claim it. A guess becomes a claim.",
        "“Who can use it” and “Anon?” are the columns that decide whether she "
        "can actually use the service. A name-and-number list quietly assumes "
        "no barriers exist.",
        "“One sentence, in her words” is shown underneath the number. Write what "
        "they do for her, not the organisation's mission statement.",
    ]:
        para(doc, line, 10, False, None, 5, style="List Bullet")

    para(doc, "", 6, space=8)
    para(doc, "Routes: contraception · youth_friendly · hiv_sti · sexual_violence "
              "· intimate_partner_violence · self_harm_risk · pregnancy_support "
              "· emotional_support     (separate several with | )",
         9.5, False, MUTED, 4)
    para(doc, "Channel: phone · sms · whatsapp · phone_whatsapp · ussd · web · walk_in",
         9.5, False, MUTED, 14)

    # --- the table -----------------------------------------------------------
    rows = []
    if CSV.exists():
        with CSV.open(encoding="utf-8-sig", newline="") as fh:
            rows = [r for r in csv.DictReader(fh) if any(v.strip() for v in r.values())]

    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    table.autofit = False

    for cell, (_, heading, width) in zip(table.rows[0].cells, COLUMNS):
        cell.width = Inches(width)
        run = cell.paragraphs[0].add_run(heading)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = PLUM

    for record in rows:
        cells = table.add_row().cells
        for cell, (key, _, width) in zip(cells, COLUMNS):
            cell.width = Inches(width)
            run = cell.paragraphs[0].add_run(record.get(key, "") or "")
            run.font.size = Pt(8.5)
            if key == "status" and (record.get(key) or "").strip() == "unverified":
                run.font.color.rgb = RUST
                run.bold = True

    for _ in range(BLANK_ROWS):
        for cell, (_, _, width) in zip(table.add_row().cells, COLUMNS):
            cell.width = Inches(width)
            cell.paragraphs[0].add_run("").font.size = Pt(8.5)

    # --- the note about the two existing rows --------------------------------
    doc.add_paragraph()
    para(doc, "About the two rows already filled in", 12, True, RUST, 4)
    para(doc,
         "Befrienders Kenya and Kenya Red Cross were carried over from the "
         "previous build. They may well be correct — but no source, date or "
         "checker was ever recorded, so they are marked unverified and the "
         "system will not surface them. If you confirm them, fill in Source, "
         "Checked by and Date, and change Status to verified.",
         10, False, None, 10)

    para(doc, "Priority order, if time is short: sexual_violence and "
              "intimate_partner_violence first, then contraception and "
              "youth_friendly, then the rest. The safeguarding routes are the "
              "ones where having nothing verified costs the most.",
         10, False, MUTED, 4)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"{OUT}  ·  {len(rows)} filled + {BLANK_ROWS} blank rows")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
